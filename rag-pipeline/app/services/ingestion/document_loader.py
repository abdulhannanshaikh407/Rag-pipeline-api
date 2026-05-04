"""
Document loader for PDF, TXT, and DOCX files.
"""
import os
import tempfile
from typing import List
import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document


class DocumentLoader:
    """
    Loads documents from disk or bytes into LangChain Document objects.
    Supports: PDF (with page metadata), TXT, DOCX.
    """

    def load_file(self, file_path: str) -> List[Document]:
        """Route to correct loader based on file extension."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext == ".txt":
            return self.load_txt(file_path)
        elif ext == ".docx":
            return self.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

    def load_pdf(self, file_path: str) -> List[Document]:
        """
        Use PyPDF2 for PDF loading.
        Each page = one Document with metadata:
          source, page_number, total_pages, file_name, file_size, created_at
        """
        reader = PdfReader(file_path)
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        created_at = os.path.getctime(file_path)

        documents = []
        for page_number, page in enumerate(reader.pages):
            text = page.extract_text()

            metadata = {
                "source": file_path,
                "page_number": page_number + 1,  # 1-indexed
                "total_pages": len(reader.pages),
                "file_name": file_name,
                "file_size": file_size,
                "created_at": created_at,
            }

            documents.append(Document(page_content=text, metadata=metadata))

        return documents

    def load_txt(self, file_path: str) -> List[Document]:
        """Load plain text. Detect encoding with chardet."""
        with open(file_path, "rb") as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result["encoding"]

        # Fallback to utf-8 if detection fails
        if encoding is None:
            encoding = "utf-8"

        with open(file_path, "r", encoding=encoding) as f:
            text = f.read()

        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        created_at = os.path.getctime(file_path)

        metadata = {
            "source": file_path,
            "page_number": 1,  # Treat as single page
            "total_pages": 1,
            "file_name": file_name,
            "file_size": file_size,
            "created_at": created_at,
        }

        return [Document(page_content=text, metadata=metadata)]

    def load_docx(self, file_path: str) -> List[Document]:
        """Use python-docx. Extract paragraphs + tables."""
        doc = DocxDocument(file_path)
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        created_at = os.path.getctime(file_path)

        full_text = []
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        text = "\n".join(full_text)

        metadata = {
            "source": file_path,
            "page_number": 1,  # Treat as single page for simplicity
            "total_pages": 1,
            "file_name": file_name,
            "file_size": file_size,
            "created_at": created_at,
        }

        return [Document(page_content=text, metadata=metadata)]

    def load_from_bytes(self, content: bytes, filename: str) -> List[Document]:
        """Save to temp file, load, then clean up."""
        # Determine extension from filename
        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        try:
            # Load the temporary file
            documents = self.load_file(tmp_file_path)
        finally:
            # Clean up the temporary file
            os.unlink(tmp_file_path)

        return documents
